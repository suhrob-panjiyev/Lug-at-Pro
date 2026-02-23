import os
import re
import sqlite3
from io import BytesIO
from pathlib import Path
from datetime import datetime

import streamlit as st
from docx import Document
from docx.shared import Pt

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Preformatted
from reportlab.lib.styles import getSampleStyleSheet

from pages.student_core import require_login
from llm_openai import generate_grammar_handout as openai_handout
from llm_gemini import generate_grammar_handout as gemini_handout

require_login()

st.set_page_config(page_title="Teacher", page_icon="👨‍🏫", layout="wide")


# =========================
# DB (Saved Materials)
# =========================
def _project_root() -> Path:
    return Path(__file__).resolve().parents[1]


DB_PATH = _project_root() / "data" / "teacher_materials.db"
DB_PATH.parent.mkdir(parents=True, exist_ok=True)


def get_conn() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL;")
    return conn


def init_teacher_db() -> None:
    conn = get_conn()
    cur = conn.cursor()
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS teacher_materials (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            topic TEXT NOT NULL,
            level TEXT NOT NULL,
            minutes INTEGER NOT NULL,
            language TEXT NOT NULL,
            provider_used TEXT NOT NULL,
            created_at TEXT NOT NULL,
            md TEXT NOT NULL
        )
        """
    )
    conn.commit()
    conn.close()


def save_material(topic: str, level: str, minutes: int, language: str, provider_used: str, md: str) -> int:
    conn = get_conn()
    cur = conn.cursor()
    cur.execute(
        """
        INSERT INTO teacher_materials (topic, level, minutes, language, provider_used, created_at, md)
        VALUES (?, ?, ?, ?, ?, ?, ?)
        """,
        (
            topic.strip(),
            level,
            int(minutes),
            language,
            provider_used,
            datetime.now().isoformat(timespec="seconds"),
            md,
        ),
    )
    conn.commit()
    mid = int(cur.lastrowid)
    conn.close()
    return mid


def list_materials(search: str = "", limit: int = 200):
    conn = get_conn()
    cur = conn.cursor()

    s = (search or "").strip()
    if s:
        like = f"%{s}%"
        cur.execute(
            """
            SELECT id, topic, level, minutes, language, provider_used, created_at
            FROM teacher_materials
            WHERE topic LIKE ? OR md LIKE ?
            ORDER BY id DESC
            LIMIT ?
            """,
            (like, like, int(limit)),
        )
    else:
        cur.execute(
            """
            SELECT id, topic, level, minutes, language, provider_used, created_at
            FROM teacher_materials
            ORDER BY id DESC
            LIMIT ?
            """,
            (int(limit),),
        )

    rows = cur.fetchall()
    conn.close()
    return rows


def get_material(mid: int):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute(
        """
        SELECT id, topic, level, minutes, language, provider_used, created_at, md
        FROM teacher_materials
        WHERE id = ?
        """,
        (int(mid),),
    )
    row = cur.fetchone()
    conn.close()
    return row


def delete_material(mid: int) -> None:
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("DELETE FROM teacher_materials WHERE id=?", (int(mid),))
    conn.commit()
    conn.close()


init_teacher_db()


# =========================
# Markdown -> Word/PDF helpers (Quick Fix)
# =========================
def _strip_md_inline(s: str) -> str:
    # removes **bold**, *italic*, `code`
    s = re.sub(r"`([^`]+)`", r"\1", s)
    s = s.replace("**", "")
    s = s.replace("*", "")
    return s


def md_to_docx_bytes(md: str) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md.splitlines()
    in_code = False
    table_buf = []

    def flush_table():
        nonlocal table_buf
        if not table_buf:
            return

        rows = []
        for r in table_buf:
            parts = [p.strip() for p in r.strip().strip("|").split("|")]
            if parts:
                rows.append(parts)

        table_buf = []
        if not rows:
            return

        # remove alignment row like |---|---|
        cleaned = []
        for r in rows:
            if all(re.match(r"^:?-{3,}:?$", c.replace(" ", "")) for c in r):
                continue
            cleaned.append(r)
        rows = cleaned
        if not rows:
            return

        cols = max(len(r) for r in rows)
        t = doc.add_table(rows=0, cols=cols)
        t.style = "Table Grid"

        for r in rows:
            row_cells = t.add_row().cells
            for i in range(cols):
                txt = _strip_md_inline(r[i]) if i < len(r) else ""
                row_cells[i].text = txt

    for line in lines:
        # Code fence
        if line.strip().startswith("```"):
            flush_table()
            in_code = not in_code
            continue

        # Collect markdown table rows
        if not in_code and line.strip().startswith("|") and line.strip().endswith("|"):
            table_buf.append(line)
            continue
        else:
            flush_table()

        if in_code:
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].font.name = "Consolas"
                p.runs[0].font.size = Pt(10)
            continue

        s = line.rstrip()
        if not s.strip():
            doc.add_paragraph("")
            continue

        # Headings
        if s.startswith("# "):
            doc.add_heading(_strip_md_inline(s[2:].strip()), level=1)
            continue
        if s.startswith("## "):
            doc.add_heading(_strip_md_inline(s[3:].strip()), level=2)
            continue
        if s.startswith("### "):
            doc.add_heading(_strip_md_inline(s[4:].strip()), level=3)
            continue

        # Bullets
        if s.lstrip().startswith(("- ", "* ")):
            doc.add_paragraph(_strip_md_inline(s.lstrip()[2:].strip()), style="List Bullet")
            continue

        # Normal
        doc.add_paragraph(_strip_md_inline(s))

    flush_table()

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def md_to_pdf_bytes(md: str, title: str = "Handout") -> bytes:
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, title=title)
    styles = getSampleStyleSheet()
    code_style = styles["Code"] if "Code" in styles else styles["BodyText"]

    story = []
    lines = md.splitlines()
    in_code = False

    for line in lines:
        if line.strip().startswith("```"):
            in_code = not in_code
            continue

        if not line.strip():
            story.append(Spacer(1, 8))
            continue

        if in_code:
            story.append(Preformatted(line, code_style))
            continue

        s = line.rstrip()
        s = _strip_md_inline(s)

        if s.startswith("# "):
            story.append(Paragraph(s[2:].strip(), styles["Heading1"]))
            continue
        if s.startswith("## "):
            story.append(Paragraph(s[3:].strip(), styles["Heading2"]))
            continue
        if s.startswith("### "):
            story.append(Paragraph(s[4:].strip(), styles["Heading3"]))
            continue

        if s.lstrip().startswith(("- ", "* ")):
            text = s.lstrip()[2:].strip()
            story.append(Paragraph("• " + text, styles["BodyText"]))
            continue

        # Basic escape for Paragraph
        esc = s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        story.append(Paragraph(esc, styles["BodyText"]))

    doc.build(story)
    return buf.getvalue()


# =========================
# Provider logic (Auto + quota-friendly)
# =========================
def _is_openai_quota_error(e: Exception) -> bool:
    msg = str(e).lower()
    return ("insufficient_quota" in msg) or ("exceeded your current quota" in msg) or ("error code: 429" in msg)


def generate_with_provider(provider_choice: str, topic: str, level: str, minutes: int, language: str):
    """
    Returns: (text, provider_used)
    provider_choice: "Auto (Gemini → OpenAI)" | "Gemini" | "OpenAI"
    """
    if provider_choice == "Gemini":
        return gemini_handout(topic=topic, level=level, minutes=minutes, language=language), "Gemini"

    if provider_choice == "OpenAI":
        # If OpenAI has no quota, this will raise
        return openai_handout(topic=topic, level=level, minutes=minutes, language=language), "OpenAI"

    # AUTO: Gemini first, then OpenAI fallback
    try:
        return gemini_handout(topic=topic, level=level, minutes=minutes, language=language), "Gemini"
    except Exception as e1:
        # If Gemini fails, try OpenAI
        try:
            return openai_handout(topic=topic, level=level, minutes=minutes, language=language), "OpenAI (fallback)"
        except Exception as e2:
            raise RuntimeError(f"Gemini error: {e1}\n\nOpenAI error: {e2}")


# =========================
# UI
# =========================
with st.sidebar:
    st.markdown("## 📘 Lug'at Pro")
    st.caption("Teacher bo‘limi")

    if st.button("🏠 Home", use_container_width=True):
        st.switch_page("app.py")
    if st.button("🎓 Student", use_container_width=True):
        st.switch_page("pages/1_Student.py")
    if st.button("👨‍🏫 Teacher", use_container_width=True):
        st.switch_page("pages/2_Teacher.py")
    if st.button("👤 Sayt haqida", use_container_width=True):
        st.switch_page("pages/3_About.py")
col_back, col_space = st.columns([1, 6])

with col_back:
    if st.button("⬅ Orqaga", use_container_width=True, key="teacher_back_home"):
        st.switch_page("app.py")

st.title("👨‍🏫 Teacher — Grammar Material Generator")

# state
if "teacher_out_md" not in st.session_state:
    st.session_state.teacher_out_md = ""
if "teacher_used_provider" not in st.session_state:
    st.session_state.teacher_used_provider = ""
if "teacher_last_saved_id" not in st.session_state:
    st.session_state.teacher_last_saved_id = None

tab_gen, tab_saved = st.tabs(["✨ Material yaratish", "📚 Saqlanganlar"])

# -------------------------
# TAB: Generate
# -------------------------
with tab_gen:
    left, right = st.columns([1, 1], gap="large")

    with left:
        st.subheader("⚙️ Sozlamalar")

        provider = st.selectbox("AI Provider", ["Auto (Gemini → OpenAI)", "Gemini", "OpenAI"], index=0)

        topic = st.text_input("Grammar mavzu", key="teacher_topic", placeholder="Present Simple")
        level = st.selectbox("Daraja", ["A1", "A2", "B1", "B2", "C1", "C2"], index=3)
        minutes = st.selectbox("Dars vaqti (minut)", [30, 45, 60, 90], index=1)
        language = st.selectbox("Tushuntirish tili", ["Uzbek", "Russian", "English"], index=0)

        save_to_library = st.checkbox("✅ Materialni saqlab qo‘yish (kutubxonaga)", value=True)

        c1, c2, c3 = st.columns(3)
        with c1:
            generate_clicked = st.button("✨ Material yaratish", type="primary", use_container_width=True)
        with c2:
            clear_clicked = st.button("🧹 Tozalash", use_container_width=True)
        with c3:
            example_clicked = st.button("⚡ Misol", use_container_width=True)

        if example_clicked:
            st.session_state.teacher_topic = "Present Perfect vs Past Simple"
            st.rerun()

        if clear_clicked:
            st.session_state.teacher_out_md = ""
            st.session_state.teacher_used_provider = ""
            st.session_state.teacher_last_saved_id = None
            st.rerun()

        st.markdown("---")
        st.caption("💡 Tavsiya: Provider = Auto qilsangiz, Gemini ishlamasa OpenAI sinab ko‘radi.")

    with right:
        st.subheader("📄 Natija")

        if st.session_state.teacher_out_md:
            md = st.session_state.teacher_out_md
            st.markdown(md)

            safe_name = (st.session_state.teacher_topic.strip().replace(" ", "_") or "grammar")

            colA, colB = st.columns(2)
            with colA:
                st.download_button(
                    "⬇️ Word (.docx)",
                    data=md_to_docx_bytes(md),
                    file_name=f"{safe_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key=f"gen_word_{safe_name}"
                )
            with colB:
                st.download_button(
                    "⬇️ PDF (.pdf)",
                    data=md_to_pdf_bytes(md, title=safe_name),
                    file_name=f"{safe_name}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    key=f"gen_pdf_{safe_name}"
                )

            if st.session_state.teacher_used_provider:
                st.caption(f"Provider: **{st.session_state.teacher_used_provider}**")
            if st.session_state.teacher_last_saved_id:
                st.caption(f"Saqlangan ID: **{st.session_state.teacher_last_saved_id}**")

        else:
            st.info("Hali material yo‘q. Chap tomonda sozlamalarni kiriting va **Material yaratish** ni bosing.")

    # Generate action
    if generate_clicked:
        t = (topic or "").strip()
        if not t:
            st.warning("Iltimos, grammar mavzuni kiriting 🙂")
            st.stop()

        with st.spinner("AI material tayyorlayapti..."):
            try:
                out, used = generate_with_provider(provider, t, level, int(minutes), language)

                st.session_state.teacher_out_md = out
                st.session_state.teacher_used_provider = used

                # Save to library
                if save_to_library:
                    mid = save_material(
                        topic=t,
                        level=level,
                        minutes=int(minutes),
                        language=language,
                        provider_used=used,
                        md=out,
                    )
                    st.session_state.teacher_last_saved_id = mid

                st.success(f"✅ Tayyor! (Provider: {used})")
                st.rerun()

            except Exception as e:
                # Special note for OpenAI quota
                if _is_openai_quota_error(e):
                    st.error("❌ OpenAI quota tugagan (429). Hozircha Gemini’dan foydalaning yoki Auto qiling.")
                    st.code(str(e))
                else:
                    st.error("❌ Generator ishlamadi. Sabab (qisqa):")
                    st.code(str(e))


# -------------------------
# TAB: Saved Materials
# -------------------------
with tab_saved:
    st.subheader("📚 Saqlangan materiallar")

    cols = st.columns([2, 1])
    with cols[0]:
        q = st.text_input("Qidirish (topic yoki matndan)", key="teacher_search", placeholder="Masalan: Present Simple")
    with cols[1]:
        refresh = st.button("🔄 Yangilash", use_container_width=True)

    rows = list_materials(search=q, limit=200)
    if not rows:
        st.info("Hali saqlangan material yo‘q. Avval material yarating va 'saqlab qo‘yish'ni yoqing.")
    else:
        options = []
        id_map = {}
        for r in rows:
            label = f"#{r['id']} | {r['topic']} | {r['level']} | {r['minutes']}m | {r['language']} | {r['created_at']}"
            options.append(label)
            id_map[label] = int(r["id"])

        selected = st.selectbox("Material tanlang", options, index=0)

        mid = id_map[selected]
        item = get_material(mid)

        if item:
            st.markdown("### 👀 Ko‘rish")
            st.caption(f"Provider: **{item['provider_used']}** | Sana: **{item['created_at']}**")
            st.markdown(item["md"])

            safe_name = (item["topic"].strip().replace(" ", "_") or f"material_{mid}")

            c1, c2, c3 = st.columns(3)
            with c1:
                st.download_button(
                    "⬇️ Word (.docx)",
                    data=md_to_docx_bytes(item["md"]),
                    file_name=f"{safe_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key=f"saved_word_{mid}"
                )
            with c2:
                st.download_button(
                    "⬇️ PDF (.pdf)",
                    data=md_to_pdf_bytes(item["md"], title=safe_name),
                    file_name=f"{safe_name}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    key=f"saved_pdf_{mid}"
                )
            with c3:
                if st.button("🗑️ O‘chirish", use_container_width=True):
                    delete_material(mid)
                    st.success("O‘chirildi ✅")
                    st.rerun()

            st.markdown("---")
            st.markdown("### ↩️ Generatorga yuklash")
            if st.button("📥 Shu materialni generatorga yuklash", type="primary", use_container_width=True):
                # load into generator tab
                st.session_state.teacher_topic = item["topic"]
                st.session_state.teacher_out_md = item["md"]
                st.session_state.teacher_used_provider = item["provider_used"]
                st.session_state.teacher_last_saved_id = item["id"]
                st.success("Generatorga yuklandi ✅")
                st.rerun()